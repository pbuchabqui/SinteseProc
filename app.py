"""
app.py — Síntese de Decisões Trabalhista
IA usada APENAS para critérios de liquidação (interpretação jurídica).
Todo o resto é regex + PyMuPDF (determinístico, rápido, sem custo).
"""

import io, json, re
import streamlit as st
import fitz
from groq import Groq
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# ── Config ────────────────────────────────────────────────────────────────────

st.set_page_config(page_title="Síntese de Decisões", page_icon="⚖️")
st.title("⚖️ Síntese de Decisões Trabalhista")
st.caption("Upload PDF → extrair → baixar Word / Markdown / Excel")

def get_secret(k):
    try: return st.secrets[k]
    except Exception:
        import os; return os.getenv(k)

GROQ_KEY = get_secret("GROQ_API_KEY")
if not GROQ_KEY:
    st.error("GROQ_API_KEY não configurada. Adicione em Settings → Secrets.")
    st.stop()

# ══════════════════════════════════════════════════════════════════════════════
# BLOCO 1 — EXTRAÇÃO DETERMINÍSTICA (sem IA)
# ══════════════════════════════════════════════════════════════════════════════

def ler_pdf(b: bytes) -> tuple[fitz.Document, str, str, list, int]:
    """Abre o PDF e retorna (doc, capa, texto_completo, toc, total_paginas).
    capa = texto da página 1 (cabeçalho PJe com partes explícitas).
    toc  = sumário estruturado do PDF (lista de [nivel, titulo, pagina]).
    """
    doc = fitz.open(stream=b, filetype="pdf")
    capa = doc[0].get_text() if len(doc) > 0 else ""
    toc  = doc.get_toc()  # sumário embutido — disponível em todos os PDFs PJe
    paginas = [f"[PÁGINA {i+1}]\n{p.get_text()}" for i, p in enumerate(doc)]
    return doc, capa, "\n".join(paginas), toc, len(paginas)


# ── 1.1 Dados do processo (regex) ────────────────────────────────────────────

def extrair_dados(txt: str, capa: str = "") -> dict:
    """
    capa = texto da página 1 do PJe (fonte mais confiável para partes).
    Partes: lidas da capa primeiro; fallback no corpo do processo.
    CPF/CNPJ: exigem label próximo para evitar falsos positivos de URLs PJe.
    Datas: buscadas nos dois sentidos (label→data e data→label).
    """
    def primeiro(padrao, texto=txt, grupo=0):
        m = re.search(padrao, texto, re.IGNORECASE)
        return m.group(grupo) if m else None

    def todos(padrao, texto=txt):
        return re.findall(padrao, texto, re.IGNORECASE)

    # ── Número do processo ──
    numero = primeiro(r"\d{7}-\d{2}\.\d{4}\.\d\.\d{2}\.\d{4}")

    # ── Vara do trabalho ──
    vara = primeiro(r"(\d+[ªº°]?\s*VARA\s+DO\s+TRABALHO[^\n]{0,60})", grupo=1) \
        or primeiro(r"(VARA\s+DO\s+TRABALHO[^\n]{0,60})", grupo=1)
    if vara: vara = vara.strip().rstrip(".,;:()")

    # ── CPF — exige label "CPF" a ≤80 chars antes, para não pegar hashes de URL PJe ──
    def formatar_cpf(d):
        return f"{d[:3]}.{d[3:6]}.{d[6:9]}-{d[9:]}" if len(d) == 11 else d

    cpf = None
    m_cpf = re.search(
        r"CPF[/\w\s\.]{0,30}?n?[º.]?\s*(\d{3}[\. ]?\d{3}[\. ]?\d{3}[\-\s]?\d{2})(?!\d)",
        txt, re.IGNORECASE
    )
    if m_cpf:
        d = re.sub(r"\D", "", m_cpf.group(1))
        cpf = formatar_cpf(d) if len(d) == 11 else None

    # ── CNPJ ──
    def formatar_cnpj(raw):
        if not raw: return None
        d = re.sub(r"\D", "", raw)
        return f"{d[:2]}.{d[2:5]}.{d[5:8]}/{d[8:12]}-{d[12:]}" if len(d) == 14 else raw

    cnpj_raw = primeiro(
        r"CNPJ[^\d]{0,20}(\d{2}[\.-]?\d{3}[\.-]?\d{3}[/]?\d{4}[-]?\d{2})", grupo=1
    ) or primeiro(r"(\d{2}[\.-]?\d{3}[\.-]?\d{3}[/]\d{4}[-]\d{2})", grupo=1)
    cnpj = formatar_cnpj(cnpj_raw)

    # ── OABs — normalizar e deduplicar ──
    def normalizar_oab(s):
        m = re.search(r"OAB[/\s]*([A-Z]{2})[/\s#nº°.]*\s*([\d\.]+)", s, re.IGNORECASE)
        if m:
            num = re.sub(r"\.", "", m.group(2))  # dígitos puros
            # Reagrupar com ponto como separador de milhar (ex: 113880 → 113.880)
            if len(num) > 3:
                num = ".".join(
                    [num[max(0, i-3):i] for i in range(len(num), 0, -3)][::-1]
                )
            return f"OAB/{m.group(1).upper()} {num}"
        return s.strip()
    oabs_raw = todos(r"OAB[/\s]*[A-Z]{2}[/\s#nº°.]*\s*[\d\.]+")
    oabs = list(dict.fromkeys(normalizar_oab(o) for o in oabs_raw))[:6]

    # ── Partes — da CAPA (página 1 PJe) ──
    # Formato canônico: RECLAMANTE: NOME / RECLAMADO: NOME / ADVOGADO: NOME
    def nome_da_capa(label, texto=None):
        src = texto or capa or txt
        m = re.search(rf"^{label}:\s*(.+)$", src, re.IGNORECASE | re.MULTILINE)
        if m:
            nome = m.group(1).strip().rstrip(".,;")
            if not re.search(r"PERITO|ADVOGADO|RECLAMANTE|RECLAMADO", nome, re.IGNORECASE):
                return nome
        return None

    reclamante = nome_da_capa("RECLAMANTE") or nome_da_capa("AUTOR[AE]?")
    reclamada  = nome_da_capa("RECLAMADO[A]?") or nome_da_capa("R[EÉ]U") \
              or nome_da_capa("EXECUTAD[AO]")

    advs_capa = re.findall(r"^ADVOGADO:\s*(.+)$", capa or "", re.IGNORECASE | re.MULTILINE)
    adv_rec   = advs_capa[0].strip() if len(advs_capa) > 0 else None
    adv_emp   = advs_capa[1].strip() if len(advs_capa) > 1 else None

    # ── Datas — busca bidirecional (label→data e data→label) ──
    def data_por_label(labels, fonte=None):
        src = fonte or txt
        for label in labels:
            # sentido normal: "Admissão: 12/06/2021" ou "Admissão ... 12/06/2021"
            m = re.search(
                rf"{label}[^:\n]{{0,30}}:\s*(\d{{2}}/\d{{2}}/\d{{4}})"
                rf"|{label}[^\d\n]{{0,60}}(\d{{2}}/\d{{2}}/\d{{4}})",
                src, re.IGNORECASE
            )
            if m: return m.group(1) or m.group(2)
            # sentido inverso: "12/06/2021 - Admissão" (formato CTPS)
            m2 = re.search(rf"(\d{{2}}/\d{{2}}/\d{{4}})\s*[-–]\s*{label}", src, re.IGNORECASE)
            if m2: return m2.group(1)
        return None

    # Ajuizamento — usar Data da Autuação da capa (mais preciso)
    data_ajuizamento = data_por_label([r"Data\s+da\s+Autua[çc][aã]o",
                                       r"Autua[çc][aã]o"], fonte=capa)                     or data_por_label([r"ajuizamento", r"distribui[çc][aã]o"])
    data_admissao  = data_por_label([r"admiss[aã]o", r"admitid[ao]"])
    data_demissao  = data_por_label([r"Rescis[aã]o\s+Contratual", r"demiss[aã]o",
                                     r"despedid[ao]", r"desligad[ao]"])

    # Função/cargo — padrão CTPS
    funcao = None
    m_func = re.search(r"Cargo\s+exercido\s+de\s+([A-Za-záéíóúâêîôûãõÁÉÍÓÚ\s\.]{3,50}?)(?:\n|$|,|\d)",
                       txt, re.IGNORECASE)
    if m_func: funcao = m_func.group(1).strip()

    return {
        "numero_processo":  numero or "Não localizado",
        "vara_trabalho":    vara,
        "reclamante":       reclamante,
        "cpf_reclamante":   cpf,
        "reclamada_1":      reclamada,
        "cnpj_reclamada_1": cnpj,
        "oabs":             oabs,
        "adv_reclamante":   adv_rec,
        "adv_reclamada_1":  adv_emp,
        "data_admissao":    data_admissao,
        "data_demissao":    data_demissao,
        "data_ajuizamento": data_ajuizamento,
        "funcao":           funcao,
    }

# ── 1.2 Localizar seções (busca bidirecional) ─────────────────────────────────

def buscar_secoes(doc: fitz.Document, toc: list, txt: str) -> dict:
    """
    Quando o TOC está disponível (PDFs PJe), cada documento é identificado
    individualmente com tipo, data e páginas exatas.
    Retorna sec["documentos_decisao"] como lista ordenada de documentos,
    além das chaves agrupadas por categoria para compatibilidade.
    """
    MAPA_TIPOS = {
        "Sentença":               "sentenca",
        "Acórdão":                "acordao",
        "Embargos de Declaração": "embargos",
        "Decisão":                "decisao",
        "Ficha Financeira":       "ficha",
        "Ficha de Registro":      "ficha",
        "Cartão de Ponto":        "ponto",
        "Espelho de Ponto":       "ponto",
    }
    DECISOES_CATS = {"sentenca", "acordao", "embargos", "decisao"}
    npags = len(doc)
    sec = {"documentos_decisao": []}

    # Títulos a excluir — não são decisões, são notificações ou petições
    EXCLUIR = re.compile(
        r"^Intima[çc][aã]o|^Manifesta[çc][aã]o|^Contrarraz|^Recurso|"
        r"^Agravo|^Certid[aã]o|^Comprovante|^Guia|^Planilha|"
        r"^Impugna[çc][aã]o|^Petição|^Procura[çc][aã]o",
        re.IGNORECASE
    )
    # Tipos de decisão relevantes para liquidação
    DECISOES_RELEVANTES = {
        "Sentença", "Acórdão", "Embargos de Declaração", "Decisão"
    }

    if toc:
        from collections import defaultdict
        grupos = defaultdict(list)

        for idx, entry in enumerate(toc):
            _, titulo, pag_ini = entry
            # Extrair o nome do tipo do título (após "N. DATA - TIPO - hash")
            titulo_limpo = re.sub(r"^\d+\.\s+\d{2}/\d{2}/\d{4}\s+-\s+", "", titulo)
            titulo_limpo = re.sub(r"\s+-\s+[a-f0-9]{7,8}$", "", titulo_limpo, flags=re.IGNORECASE)

            # Excluir intimações, petições e outros documentos não-decisórios
            if EXCLUIR.search(titulo_limpo):
                continue

            tipo_cat = next(
                (cat for tipo_str, cat in MAPA_TIPOS.items()
                 if tipo_str.lower() in titulo_limpo.lower()),
                None
            )
            if not tipo_cat: continue

            pag_fim = (toc[idx+1][2] - 1) if idx+1 < len(toc) else npags
            p0, p1 = pag_ini - 1, pag_fim - 1  # 0-based
            texto_doc = "\n".join(doc[i].get_text() for i in range(p0, min(p1+1, npags)))

            # Data do título do TOC (confiável)
            data_m = re.search(r"(\d{2}/\d{2}/\d{4})", titulo)
            data = data_m.group(1) if data_m else None

            if tipo_cat in DECISOES_CATS:
                # Embargos: excluir petições da parte (sem assinatura de juiz)
                if tipo_cat == "embargos":
                    if not re.search(r"Juiz\b|Desembargador|Ministro", texto_doc, re.IGNORECASE):
                        continue  # petição da parte, não decisão judicial

                # Guardar documento individual — usado por extrair_decisoes
                sec["documentos_decisao"].append({
                    "tipo_cat": tipo_cat,
                    "tipo_label": {
                        "sentenca": "Sentença",
                        "acordao":  "Acórdão",
                        "embargos": "Embargos de Declaração",
                        "decisao":  "Decisão",
                    }.get(tipo_cat, tipo_cat.capitalize()),
                    "titulo":   titulo,
                    "data":     data,
                    "texto":    texto_doc,
                })
                grupos[tipo_cat].append(f"\n--- {titulo} ---\n{texto_doc}")
            elif tipo_cat in ("ficha", "ponto"):
                grupos[tipo_cat].append(texto_doc)

        for cat, blocos in grupos.items():
            sec[cat] = "\n\n".join(blocos)

    # Fallback textual para PDFs sem TOC
    if not sec["documentos_decisao"]:
        linhas = txt.split("\n")
        padroes = {
            "sentenca": r"VISTOS[,\s]+RELATADOS|VISTOS[,\s]+ETC",
            "acordao":  r"A\s*C\s*[OÓ]\s*R\s*D\s*[AÃ]\s*O",
        }
        for nome, p in padroes.items():
            if nome in sec: continue
            for i in range(len(linhas)-1, -1, -1):
                if re.search(p, linhas[i], re.IGNORECASE):
                    sec[nome] = "\n".join(linhas[i:i+600])
                    break

    # Ficha e ponto via busca textual se não vieram do TOC
    for nome, padrao in [
        ("ficha", r"FICHA FINANCEIRA|CONTRACHEQUE|HOLERITE"),
        ("ponto", r"CART[AÃ]O DE PONTO|ESPELHO DE PONTO"),
    ]:
        if not sec.get(nome):
            linhas = txt.split("\n")
            for i, l in enumerate(linhas):
                if re.search(padrao, l, re.IGNORECASE):
                    sec[nome] = "\n".join(linhas[i:i+600]); break

    return sec
# ── 1.3 Extrair decisões (sem IA — extração literal de texto) ─────────────────

MARCADORES_TIPO = {
    "Sentença":              r"S\s*E\s*N\s*T\s*E\s*N\s*[CÇ]\s*A|VISTOS[,\s]+RELATADOS",
    "Acórdão":               r"A\s*C\s*[OÓ]\s*R\s*D\s*[AÃ]\s*O",
    "Embargos de Declaração":r"EMBARGOS\s+DE\s+DECLARA[CÇ][AÃ]O",
}
MARCADORES_DISPOSITIVO = [
    r"ISSO\s+POSTO", r"ISTO\s+POSTO", r"ANTE\s+O\s+EXPOSTO",
    r"DIANTE\s+DO\s+EXPOSTO", r"PELO\s+EXPOSTO",
    r"DECIDO\s*[:;]", r"DECIDE-SE", r"JULGO",
    r"ACORDAM\s+os",       # acórdão TRT
    r"^(III\s*\)\s*)?CONCLUSÃO",  # decisões TST/TRT monocráticas
]
MARCADORES_FIM_DISPOSITIVO = [
    r"Intimem-se", r"Cumpra-se", r"Publique-se", r"Registre-se",
    r"Após\s+o\s+trânsito", r"P\.R\.I", r"Intime-se",
]
VERBAS_CONHECIDAS = [
    "horas extras", "adicional noturno", "adicional de periculosidade",
    "adicional de insalubridade", "férias", "13º salário", "aviso prévio",
    "FGTS", "multa do art. 477", "multa do art. 467", "dano moral",
    "dano material", "diferenças salariais", "equiparação salarial",
    "intervalo intrajornada", "DSR", "vale-transporte", "vale-alimentação",
    "horas in itinere", "sobreaviso", "prontidão",
]

def extrair_decisao_de_doc(doc_info: dict) -> dict:
    """
    Extrai dados de uma decisão a partir de um documento individual do TOC.
    O tipo e a data já são conhecidos — só precisa localizar o dispositivo.
    """
    tipo_label = doc_info["tipo_label"]
    data       = doc_info["data"]  # data confiável do TOC
    texto      = doc_info["texto"]
    linhas     = texto.split("\n")
    bloco      = texto

    # ── Dispositivo ──────────────────────────────────────────────────────────────
    # Para sentenças longas, priorizar marcadores condenatórios sobre os introdutórios
    # Ordem de preferência: Ante o exposto > Condeno > ISSO POSTO > ...
    MARCADORES_CONDENATORIO = [
        r"Ante\s+o\s+exposto[,\.]",          # sentença TRT — resume a condenação
        r"(?:^|\n)\s*Condeno\b",              # condenação direta
        r"(?:^|\n)\s*JULGO\s+PROCED",
        r"NEGO\s+PROVIMENTO",                 # acórdão — nega o recurso
        r"DAR\s+PARCIAL\s+PROVIMENTO",        # acórdão — provimento parcial
        r"DAR\s+PROVIMENTO",                  # acórdão — provimento total
        r"ACORDAM\s+os",                      # acórdão — início do decisum
        r"Nego\s+seguimento",                 # TST/TRT — inadmite recurso
        r"denego\s+seguimento",               # TST — nega AIRR
        r"^(III\s*\)\s*)?CONCLUSÃO",          # TST monocrático
        r"Recebo\s+o\s+recurso",              # despacho de admissão
    ]
    MARCADORES_INTRO = [                      # usados só se nada acima for encontrado
        r"ISSO\s+POSTO", r"ISTO\s+POSTO",
        r"DIANTE\s+DO\s+EXPOSTO", r"PELO\s+EXPOSTO",
        r"DECIDO\s*[:;]", r"DECIDE-SE",
    ]

    dispositivo = ""
    inicio_disp = None

    # 1ª passagem: marcadores condenatórios (mais precisos)
    for j, l in enumerate(linhas):
        if any(re.search(p, l, re.IGNORECASE | re.MULTILINE) for p in MARCADORES_CONDENATORIO):
            inicio_disp = j
            break

    # 2ª passagem: marcadores introdutórios (fallback)
    if inicio_disp is None:
        for j, l in enumerate(linhas):
            if any(re.search(p, l, re.IGNORECASE) for p in MARCADORES_INTRO):
                inicio_disp = j
                break

    if inicio_disp is not None:
        linhas_disp = []
        for l in linhas[inicio_disp:]:
            linhas_disp.append(l)
            if any(re.search(p, l, re.IGNORECASE) for p in MARCADORES_FIM_DISPOSITIVO):
                break
        dispositivo = "\n".join(linhas_disp).strip()
    else:
        linhas_uteis = [l for l in linhas if l.strip() and not re.search(
            r"^Fls\.|^Documento assinado|^https?://|^Número do (processo|documento)|^Certificado",
            l.strip()
        )]
        dispositivo = "\n".join(linhas_uteis).strip()

    # ── Resultado — baseado no tipo e conteúdo do documento ──────────────────────
    tipo_cat = doc_info.get("tipo_cat", "")
    titulo   = doc_info.get("titulo", "").lower()

    if tipo_cat == "decisao":
        if re.search(r"Recebo\s+o\s+recurso|recurso.*recebido", bloco, re.IGNORECASE):
            resultado = "despacho — recurso admitido"
        elif re.search(r"denego\s+seguimento|negado\s+seguimento|intranscend", bloco, re.IGNORECASE):
            resultado = "recurso não admitido (TST)"  # verifica antes de "Nego seguimento"
        elif re.search(r"Nego\s+seguimento|não\s+admito|deserção", bloco, re.IGNORECASE):
            resultado = "recurso não admitido"
        else:
            resultado = "decisão interlocutória"

    elif tipo_cat == "embargos":
        if re.search(r"NEGO\s+PROVIMENTO|embargos.*rejeit|não\s+merece\s+prosperar", bloco, re.IGNORECASE):
            resultado = "embargos rejeitados"
        elif re.search(r"ACOLHO|acolhidos", bloco, re.IGNORECASE):
            resultado = "embargos acolhidos"
        else:
            resultado = "embargos parcialmente acolhidos"

    elif tipo_cat == "acordao":
        if re.search(r"DAR\s+PARCIAL\s+PROVIMENTO", bloco, re.IGNORECASE):
            resultado = "parcialmente provido"
        elif re.search(r"DAR\s+PROVIMENTO\b", bloco, re.IGNORECASE):
            resultado = "provido"
        elif re.search(r"NEGAR\s+PROVIMENTO|NEGO\s+PROVIMENTO|unanimidade.*negar", bloco, re.IGNORECASE):
            resultado = "negado provimento"
        else:
            resultado = "parcialmente provido"

    else:  # sentenca — inclui julgamento de embargos de declaração
        # Sentença sobre embargos: tipo_cat é "sentenca" mas trata embargos
        if re.search(r"embargos\s+de\s+declara", bloco, re.IGNORECASE) and \
           re.search(r"NEGO\s+PROVIMENTO|não\s+merece\s+prosperar|rejeito\s+os\s+embargos", bloco, re.IGNORECASE):
            resultado = "embargos rejeitados"
        elif re.search(r"embargos\s+de\s+declara", bloco, re.IGNORECASE) and \
             re.search(r"ACOLHO|acolhidos", bloco, re.IGNORECASE):
            resultado = "embargos acolhidos"
        elif re.search(r"JULGO\s+IMPROCEDENTE|totalmente\s+improcedente", bloco, re.IGNORECASE):
            resultado = "improcedente"
        elif re.search(r"PROCEDENTE[^S]\b|integralmente\s+procedente", bloco, re.IGNORECASE):
            resultado = "procedente"
        else:
            resultado = "parcialmente procedente"

    # Verbas — lista pré-definida
    verbas = [v for v in VERBAS_CONHECIDAS
              if re.search(re.escape(v), bloco, re.IGNORECASE)]

    return {
        "tipo":                 tipo_label,
        "data":                 data,
        "dispositivo":          dispositivo or "(dispositivo não localizado)",
        "resultado_reclamante": resultado,
        "verbas_deferidas":     verbas,
    }


def extrair_decisoes(secs: dict) -> list[dict]:
    """
    Extrai todas as decisões do processo.
    Quando o TOC está disponível (secs["documentos_decisao"]),
    processa cada documento individualmente — tipo e data já são conhecidos.
    Fallback: busca textual no texto concatenado.
    """
    documentos = secs.get("documentos_decisao", [])

    if documentos:
        # TOC disponível — processar cada documento individualmente
        return [extrair_decisao_de_doc(d) for d in documentos]

    # Fallback: texto concatenado (PDFs sem TOC)
    txt = "\n\n".join(filter(None, [
        secs.get("sentenca"), secs.get("acordao"),
        secs.get("embargos"), secs.get("decisao"),
    ]))
    if not txt:
        return []

    decisoes = []
    linhas = txt.split("\n")
    i = 0
    while i < len(linhas):
        linha = linhas[i]
        tipo_encontrado = None
        for tipo, padrao in MARCADORES_TIPO.items():
            if re.search(padrao, linha, re.IGNORECASE):
                tipo_encontrado = tipo
                break
        if not tipo_encontrado:
            i += 1; continue

        janela = linhas[i:i+600]
        bloco  = "\n".join(janela)
        m_data = re.search(r"\d{2}/\d{2}/\d{4}", bloco)
        data   = m_data.group() if m_data else None

        inicio_disp = None
        for j, l in enumerate(janela):
            if any(re.search(p, l, re.IGNORECASE) for p in MARCADORES_DISPOSITIVO):
                inicio_disp = j; break

        dispositivo = ""
        if inicio_disp is not None:
            linhas_disp = []
            for l in janela[inicio_disp:]:
                linhas_disp.append(l)
                if any(re.search(p, l, re.IGNORECASE) for p in MARCADORES_FIM_DISPOSITIVO):
                    break
            dispositivo = "\n".join(linhas_disp).strip()

        resultado = "parcialmente procedente"
        if re.search(r"JULGO\s+IMPROCEDENTE|NEGO\s+PROVIMENTO", bloco, re.IGNORECASE):
            resultado = "improcedente / negado provimento"
        elif re.search(r"DAR\s+PARCIAL\s+PROVIMENTO", bloco, re.IGNORECASE):
            resultado = "parcialmente provido"

        verbas = [v for v in VERBAS_CONHECIDAS if re.search(re.escape(v), bloco, re.IGNORECASE)]

        decisoes.append({
            "tipo": tipo_encontrado, "data": data,
            "dispositivo": dispositivo or "(dispositivo não localizado)",
            "resultado_reclamante": resultado, "verbas_deferidas": verbas,
        })
        i += 400

    return decisoes


# ── 1.4 Ficha financeira (PyMuPDF tabelas — sem IA) ───────────────────────────

def extrair_ficha(doc: fitz.Document) -> dict:
    """
    Extrai ficha financeira usando detecção de tabelas do PyMuPDF.
    Sem IA — dados tabulares são determinísticos.
    """
    rubricas_set = set()
    competencias = []

    for pagina in doc:
        texto_pag = pagina.get_text()
        # Verificar se esta página tem ficha financeira
        if not re.search(
            r"FICHA FINANCEIRA|CONTRACHEQUE|HOLERITE|FOLHA DE PAGAMENTO",
            texto_pag, re.IGNORECASE
        ):
            continue

        try:
            tabelas = pagina.find_tables()
        except Exception:
            continue

        for tabela in tabelas:
            try:
                linhas = tabela.extract()
            except Exception:
                continue
            if not linhas or len(linhas) < 2:
                continue

            cabecalho = [str(c).strip() if c else "" for c in linhas[0]]

            # Detectar coluna de competência
            col_comp = next(
                (i for i, c in enumerate(cabecalho)
                 if re.search(r"compet[êe]ncia|m[êe]s|per[íi]odo", c, re.IGNORECASE)),
                0
            )

            for linha in linhas[1:]:
                if not linha or not any(linha):
                    continue
                comp_val = str(linha[col_comp]).strip() if linha[col_comp] else ""
                if not re.search(r"\d{2}/\d{4}|\d{4}", comp_val):
                    continue

                valores = {}
                for ci, cel in enumerate(linha):
                    if ci == col_comp: continue
                    nome_col = cabecalho[ci] if ci < len(cabecalho) else f"col_{ci}"
                    if not nome_col: continue
                    val_str = str(cel).strip() if cel else ""
                    # Tentar converter para float
                    val_num = None
                    try:
                        val_num = float(
                            val_str.replace(".", "").replace(",", ".").replace("R$","").strip()
                        )
                    except (ValueError, AttributeError):
                        pass
                    if val_str:
                        rubricas_set.add(nome_col)
                        valores[nome_col] = {"referencia": "", "valor": val_num or val_str}

                if valores:
                    competencias.append({
                        "competencia": comp_val,
                        "valores": valores,
                        "total_proventos": None,
                        "inss_base": None,
                        "inss_desconto": None,
                        "fgts_base": None,
                    })

    if not competencias:
        return {"erro": "Ficha não localizada ou sem tabelas detectáveis no PDF"}

    return {"rubricas": sorted(rubricas_set), "competencias": competencias}


# ── 1.5 Espelho de ponto (PyMuPDF tabelas — sem IA) ───────────────────────────

def extrair_ponto(doc: fitz.Document) -> dict:
    """
    Extrai espelho de ponto usando detecção de tabelas do PyMuPDF.
    Sem IA — dados tabulares são determinísticos.
    """
    registros = []
    DIAS = {"SEG","TER","QUA","QUI","SEX","SÁB","SAB","DOM"}

    for pagina in doc:
        texto_pag = pagina.get_text()
        if not re.search(
            r"CART[AÃ]O DE PONTO|ESPELHO DE PONTO|REGISTRO DE PONTO",
            texto_pag, re.IGNORECASE
        ):
            continue

        try:
            tabelas = pagina.find_tables()
        except Exception:
            continue

        for tabela in tabelas:
            try:
                linhas = tabela.extract()
            except Exception:
                continue
            if not linhas or len(linhas) < 2:
                continue

            for linha in linhas[1:]:
                if not linha: continue
                celulas = [str(c).strip() if c else "" for c in linha]

                # Identificar célula de data
                data = next(
                    (c for c in celulas if re.match(r"\d{2}/\d{2}/\d{4}", c)),
                    None
                )
                if not data: continue

                # Dia da semana
                dia = next((c for c in celulas if c.upper() in DIAS), "")

                # Horários (HH:MM)
                horarios = [c for c in celulas if re.match(r"\d{2}:\d{2}", c)]

                # Horas trabalhadas e extras (último par de HH:MM costuma ser totais)
                ht = horarios[-2] if len(horarios) >= 2 else ""
                he = horarios[-1] if len(horarios) >= 1 else ""

                # Observação
                obs = next(
                    (c for c in celulas
                     if re.search(r"DSR|FALTA|FERIADO|FOLGA|AFASTAMENTO|FÉRIAS", c, re.IGNORECASE)),
                    ""
                )

                registros.append({
                    "data":             data,
                    "dia_semana":       dia,
                    "entradas_saidas":  horarios[:-2] if len(horarios) > 2 else horarios,
                    "horas_trabalhadas":ht,
                    "horas_extras":     he,
                    "observacao":       obs,
                })

    if not registros:
        return {"erro": "Ponto não localizado ou sem tabelas detectáveis no PDF"}

    return {"registros": registros}


# ══════════════════════════════════════════════════════════════════════════════
# BLOCO 2 — IA (apenas critérios de liquidação)
# ══════════════════════════════════════════════════════════════════════════════

MODELO_RACIOCINIO = "openai/gpt-oss-120b"

# ── Base de conhecimento trabalhista (de references/trabalhista.md do ContextAI) ──
CONHECIMENTO_TRABALHISTA = """
## Critérios de liquidação frequentes

### Horas extras — bancários
- Divisor 180: bancário comum (art. 224, caput, CLT).
- Divisor 220: cargo de confiança (art. 224, §2º, CLT).
- Base de cálculo: salário + todas as verbas de natureza salarial — Súmula 264/TST.
- Adicional: o percentual convencional (CCT) prevalece sobre o legal.

### Reflexos (incidências)
RSR/sábados/feriados | férias + 1/3 | 13º salário | FGTS (sem multa 40% se contrato vigente) | aviso prévio (apenas em rescisão).

### Tema 9/TST — IRR 0010169-57.2013.5.05.0024
- RSR majorado por horas extras repercute em férias, 13º e FGTS APENAS a partir de 20/03/2023.
- Período anterior: OJ 394/SBDI-1/TST (não repercute).
- Sempre dividir o cálculo nesses dois marcos temporais.

### Atualização monetária (pós-ADC 58/STF)
- IPCA-E até 31/12/2021.
- SELIC simples a partir de 01/01/2022 (juros + correção unificados).
- NÃO duplicar juros e correção no mesmo período.

### Jurisprudência de referência
- Bancário cargo de confiança: Art. 224 §2º CLT; Súmula 102/TST
- Gratificação de função: Súmula 109/TST
- Divisor bancário: Súmula 124/TST
- Base de cálculo HE: Súmula 264/TST
- RSR e HE: OJ 394/SBDI-1; IRR Tema 9/TST
- Atualização monetária: ADC 58/STF; Súmula 381/TST
- INSS sobre parcelas: Súmula 368/TST (apurado mês a mês sobre valor histórico)
- IRPF: Art. 12-A Lei 7.713/1988 (tabela progressiva acumulada)

### INSS — Súmula 26/TRT-4
Descontos apurados mês a mês sobre o valor histórico, com exclusão dos juros de mora,
respeitado o limite máximo mensal do salário de contribuição, observadas as alíquotas
vigentes à época e os valores já recolhidos.
"""

# ── Regras absolutas de análise pericial (de instrucoes-analise.md do ContextAI) ──
REGRAS_ANALISE = """
## Regras absolutas
- Responda sempre em português brasileiro formal.
- Seja preciso com valores monetários e datas (use vírgula decimal, R$).
- Se o contexto for insuficiente, diga qual informação falta — NUNCA invente dados.
- Baseie-se apenas no conteúdo fornecido do processo.
- Cite súmulas, OJs, artigos de lei e teses quando aplicáveis.
- Documentação incompleta: diga claramente e oriente como proceder.
"""


def chamar_ia(txt: str, instrucao: str, limite: int = 30000,
              incluir_base_trabalhista: bool = False) -> dict:
    """Chamada ao Groq com retry automático em caso de rate limit."""
    import time

    modelo = get_secret("GROQ_MODEL") or "llama-3.3-70b-versatile"
    usa_raciocinio = modelo == MODELO_RACIOCINIO

    base = CONHECIMENTO_TRABALHISTA if incluir_base_trabalhista else ""
    params = dict(
        model=modelo, stream=False,
        messages=[
            {"role":"system","content":(
                "Você é perito contábil trabalhista brasileiro especializado em liquidação de sentença no TRT4.\n"
                + REGRAS_ANALISE
                + (f"\n\n## Base de referência jurídica\n{base}" if base else "")
                + "\n\nResponda APENAS com JSON válido, sem texto antes/depois, sem ```json```.")},
            {"role":"user","content":f"{instrucao}\n\nTEXTO:\n{txt[:limite]}"},
        ]
    )
    if usa_raciocinio:
        params.update({"max_completion_tokens":3000,"temperature":0.6,"top_p":1,
                       "reasoning_effort": get_secret("GROQ_REASONING_EFFORT") or "medium"})
    else:
        params.update({"max_tokens":3000,"temperature":0.1})

    for tentativa in range(3):
        try:
            r   = Groq(api_key=GROQ_KEY).chat.completions.create(**params)
            raw = r.choices[0].message.content
            try:
                return json.loads(re.sub(r"```(?:json)?|```","",raw).strip())
            except json.JSONDecodeError:
                return {"erro": "JSON inválido", "texto_bruto": raw}

        except Exception as e:
            msg = str(e)
            if "rate_limit" in msg.lower() or "429" in msg:
                if tentativa < 2:
                    espera = 65 if tentativa == 0 else 90
                    st.warning(f"⏳ Rate limit — aguardando {espera}s (tentativa {tentativa+1}/3)...")
                    time.sleep(espera)
                    continue
                return {"erro": "Rate limit persistente. Tente novamente em alguns minutos."}
            return {"erro": f"Erro na API: {msg}"}

    return {"erro": "Falha após 3 tentativas."}


def ext_criterios(dispositivo: str) -> dict:
    """
    Única chamada de IA no sistema.
    Interpreta o dispositivo da sentença para extrair parâmetros de liquidação.
    Isso exige compreensão semântica — regex não resolve.
    """
    return chamar_ia(dispositivo, """
Analise o dispositivo da sentença/acórdão trabalhista e extraia os critérios de liquidação.
Use a base de referência jurídica fornecida para preencher critérios não explicitados na sentença
(ex: se não há índice de correção, aplicar SELIC por ADC 58; se não há divisor, usar o padrão da categoria).

{
  "criterios": {
    "base_salarial": "descrição da base (ex: última remuneração R$ X)",
    "periodo_apurado": {"inicio": "dd/mm/aaaa", "fim": "dd/mm/aaaa"},
    "jornada_contratual": "ex: 8h diárias / 44h semanais",
    "jornada_real_apurada": "ex: 10h diárias conforme cartões de ponto",
    "divisor": "220 ou 180 ou outro — justificar",
    "adicional_horas_extras": "50% ou 100% ou percentual CCT",
    "reflexos": ["DSR", "férias", "13º salário", "aviso prévio", "FGTS"],
    "marco_tema9": "aplicar Tema 9/TST a partir de 20/03/2023 se houver RSR majorado",
    "fgts": {"base": "todas as verbas deferidas", "multa_40": true},
    "atualizacao_monetaria": "IPCA-E até 31/12/2021 + SELIC a partir de 01/01/2022 (ADC 58)",
    "juros": "incluídos na SELIC a partir de 01/01/2022",
    "inss_empregado": "Súmula 26/TRT-4: mês a mês sobre valor histórico",
    "inss_patronal": "a cargo da reclamada",
    "ir": "tabela progressiva acumulada (art. 12-A Lei 7.713/1988) ou isento",
    "exclusoes_expressas": ["ex: dano moral não integra base FGTS"],
    "observacoes": "outros critérios relevantes incluindo verbas deferidas e deduções autorizadas"
  }
}
""", limite=30000, incluir_base_trabalhista=True)


def ext_alertas_periciais(texto_processo: str) -> dict:
    """
    Gera a Seção 8 — Alertas Periciais (obrigatória conforme instrucoes-analise.md).
    Usa o texto das decisões + despacho de nomeação para identificar riscos e atenções.
    """
    return chamar_ia(texto_processo, """
Analise o processo trabalhista e gere os ALERTAS PERICIAIS para o perito.
Esta é a seção mais importante do relatório — seja específico e prático.

{
  "alertas": {
    "formato_laudo": "PJe-Calc obrigatório | laudo livre | verificar despacho",
    "prazo_laudo": "data ou 'verificar despacho de nomeação'",
    "pje_calc_exigido": true,
    "documentacao_faltante": ["lista do que não está nos autos mas é necessário"],
    "pontos_atencao": [
      "alertas específicos do caso — riscos, armadilhas, divergências prováveis"
    ],
    "marcos_temporais_criticos": [
      "ex: contrato anterior a 20/03/2023 — verificar Tema 9/TST para RSR"
    ],
    "vincendas": "apurar parcelas vincendas até data-base do laudo (art. 899/CLT)",
    "honorarios_risco": "OJ 19/TRT-3: divergência significativa pode gerar condenação em honorários",
    "observacoes_finais": "outros alertas não cobertos acima"
  }
}
""", limite=20000, incluir_base_trabalhista=True)

def gerar_word(dados, decisoes, criterios, alertas=None) -> bytes:
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)

    def h(t, n=1):
        x = doc.add_heading(t, level=n)
        x.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def campo(label, val):
        if not val: return
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(str(val))

    h("SÍNTESE DE DECISÕES PARA LIQUIDAÇÃO")

    h("1. Identificação do Processo", 2)
    campo("Número",    dados.get("numero_processo"))
    campo("Vara",      dados.get("vara_trabalho"))
    campo("Reclamante",dados.get("reclamante"))
    campo("CPF",       dados.get("cpf_reclamante"))
    campo("Reclamada", dados.get("reclamada_1"))
    campo("CNPJ",      dados.get("cnpj_reclamada_1"))
    campo("Advogado (reclamante)", dados.get("adv_reclamante"))
    campo("Advogado (reclamada)",  dados.get("adv_reclamada_1"))
    campo("Admissão",    dados.get("data_admissao"))
    campo("Demissão",    dados.get("data_demissao"))
    campo("Ajuizamento", dados.get("data_ajuizamento"))
    if dados.get("oabs"):
        campo("OABs encontradas", " | ".join(dados["oabs"]))

    if decisoes:
        h("2. Decisões Judiciais", 2)
        for i, d in enumerate(decisoes, 1):
            h(f"2.{i} {d.get('tipo','Decisão')}", 3)
            campo("Data",      d.get("data"))
            campo("Resultado", d.get("resultado_reclamante"))
            if d.get("verbas_deferidas"):
                campo("Verbas identificadas", ", ".join(d["verbas_deferidas"]))
            doc.add_paragraph().add_run("DISPOSITIVO:").bold = True
            doc.add_paragraph(d.get("dispositivo",""))

    if criterios and "erro" not in criterios and criterios.get("criterios"):
        h("3. Critérios de Liquidação", 2)
        c = criterios["criterios"]
        per = c.get("periodo_apurado",{})
        if per: campo("Período", f"{per.get('inicio')} a {per.get('fim')}")
        for lb, ch in [
            ("Base salarial","base_salarial"),
            ("Jornada contratual","jornada_contratual"),
            ("Jornada real apurada","jornada_real_apurada"),
            ("Divisor","divisor"),
            ("Adicional HE","adicional_horas_extras"),
            ("Atualização monetária","atualizacao_monetaria"),
            ("Juros","juros"),
            ("INSS empregado","inss_empregado"),
            ("INSS patronal","inss_patronal"),
            ("IR","ir"),
            ("Observações","observacoes"),
        ]:
            campo(lb, c.get(ch))
        if c.get("reflexos"):
            campo("Reflexos", ", ".join(c["reflexos"]))
        fgts = c.get("fgts",{})
        if fgts:
            campo("FGTS", f"{fgts.get('base','')} — {'com' if fgts.get('multa_40') else 'sem'} multa 40%")
        if c.get("exclusoes_expressas"):
            campo("Exclusões expressas", "; ".join(c["exclusoes_expressas"]))
        if c.get("marco_tema9"):
            campo("Marco Tema 9/TST", c["marco_tema9"])

    # Seção 4 — Alertas Periciais (obrigatório conforme instrucoes-analise.md)
    if alertas and "erro" not in alertas and alertas.get("alertas"):
        a = alertas["alertas"]
        h("4. Alertas Periciais", 2)
        p = doc.add_paragraph()
        run = p.add_run("⚠️ Esta seção deve ser lida antes de iniciar o laudo.")
        run.bold = True; run.font.color.rgb = None

        def alerta(label, val):
            if not val: return
            if isinstance(val, list) and not val: return
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{label}: ").bold = True
            p.add_run(", ".join(val) if isinstance(val, list) else str(val))

        campo("Formato do laudo",    a.get("formato_laudo"))
        campo("Prazo",               a.get("prazo_laudo"))
        campo("PJe-Calc exigido",    "Sim" if a.get("pje_calc_exigido") else "Verificar")
        campo("Apuração vincendas",  a.get("vincendas"))
        campo("Risco honorários",    a.get("honorarios_risco"))

        if a.get("documentacao_faltante"):
            h("Documentação faltante", 3)
            for item in a["documentacao_faltante"]:
                doc.add_paragraph(item, style="List Bullet")

        if a.get("pontos_atencao"):
            h("Pontos de atenção", 3)
            for item in a["pontos_atencao"]:
                doc.add_paragraph(item, style="List Bullet")

        if a.get("marcos_temporais_criticos"):
            h("Marcos temporais críticos", 3)
            for item in a["marcos_temporais_criticos"]:
                doc.add_paragraph(item, style="List Bullet")

        campo("Observações", a.get("observacoes_finais"))

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue()


def gerar_markdown(dados, decisoes, criterios, alertas=None) -> str:
    md = ["# SÍNTESE DE DECISÕES PARA LIQUIDAÇÃO\n",
          "## 1. Identificação do Processo\n"]

    for lb, ch in [
        ("Número","numero_processo"),("Vara","vara_trabalho"),
        ("Reclamante","reclamante"),("CPF","cpf_reclamante"),
        ("Reclamada","reclamada_1"),("CNPJ","cnpj_reclamada_1"),
        ("Advogado reclamante","adv_reclamante"),
        ("Advogado reclamada","adv_reclamada_1"),
        ("Admissão","data_admissao"),("Demissão","data_demissao"),
        ("Ajuizamento","data_ajuizamento"),
    ]:
        if dados.get(ch): md.append(f"**{lb}:** {dados[ch]}")
    if dados.get("oabs"):
        md.append(f"**OABs:** {' | '.join(dados['oabs'])}")

    if decisoes:
        md.append("\n## 2. Decisões Judiciais\n")
        for i, d in enumerate(decisoes, 1):
            md.append(f"### 2.{i} {d.get('tipo','Decisão')}")
            md.append(f"**Data:** {d.get('data','')}  |  **Resultado:** {d.get('resultado_reclamante','')}")
            if d.get("verbas_deferidas"):
                md.append(f"**Verbas:** {', '.join(d['verbas_deferidas'])}")
            md.append(f"\n**DISPOSITIVO:**\n\n{d.get('dispositivo','')}\n")

    if criterios and "erro" not in criterios and criterios.get("criterios"):
        md.append("\n## 3. Critérios de Liquidação\n")
        c = criterios["criterios"]
        per = c.get("periodo_apurado",{})
        if per: md.append(f"**Período:** {per.get('inicio')} a {per.get('fim')}")
        for lb, ch in [
            ("Base salarial","base_salarial"),("Jornada contratual","jornada_contratual"),
            ("Jornada real","jornada_real_apurada"),("Divisor","divisor"),
            ("Adicional HE","adicional_horas_extras"),("Atualização","atualizacao_monetaria"),
            ("Juros","juros"),("INSS empregado","inss_empregado"),
            ("INSS patronal","inss_patronal"),("IR","ir"),("Observações","observacoes"),
        ]:
            if c.get(ch): md.append(f"**{lb}:** {c[ch]}")
        if c.get("reflexos"): md.append(f"**Reflexos:** {', '.join(c['reflexos'])}")
        fgts = c.get("fgts",{})
        if fgts: md.append(f"**FGTS:** {fgts.get('base','')} — {'com' if fgts.get('multa_40') else 'sem'} multa 40%")
        if c.get("exclusoes_expressas"):
            md.append(f"**Exclusões:** {'; '.join(c['exclusoes_expressas'])}")
        if c.get("marco_tema9"):
            md.append(f"**Marco Tema 9/TST:** {c['marco_tema9']}")

    if alertas and "erro" not in alertas and alertas.get("alertas"):
        a = alertas["alertas"]
        md.append("\n## 4. Alertas Periciais ⚠️\n")
        if a.get("formato_laudo"):   md.append(f"**Formato:** {a['formato_laudo']}")
        if a.get("prazo_laudo"):     md.append(f"**Prazo:** {a['prazo_laudo']}")
        if a.get("pje_calc_exigido") is not None:
            md.append(f"**PJe-Calc:** {'Exigido' if a['pje_calc_exigido'] else 'Verificar'}")
        if a.get("vincendas"):       md.append(f"**Vincendas:** {a['vincendas']}")
        if a.get("honorarios_risco"):md.append(f"**Risco honorários:** {a['honorarios_risco']}")
        if a.get("documentacao_faltante"):
            md.append("\n### Documentação faltante")
            for item in a["documentacao_faltante"]: md.append(f"- {item}")
        if a.get("pontos_atencao"):
            md.append("\n### Pontos de atenção")
            for item in a["pontos_atencao"]: md.append(f"- {item}")
        if a.get("marcos_temporais_criticos"):
            md.append("\n### Marcos temporais críticos")
            for item in a["marcos_temporais_criticos"]: md.append(f"- {item}")
        if a.get("observacoes_finais"): md.append(f"\n**Observações:** {a['observacoes_finais']}")

    return "\n\n".join(md)


AZUL, BRANCO = "1F497D", "FFFFFF"

def _cab(c):
    c.font = Font(bold=True, color=BRANCO, size=11)
    c.fill = PatternFill("solid", fgColor=AZUL)
    c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

def _borda(c):
    b = Side(style="thin")
    c.border = Border(left=b, right=b, top=b, bottom=b)

def _campo_xl(ws, ln, label, val):
    ca = ws.cell(ln,1,label); cb = ws.cell(ln,2,val or "")
    ca.font = Font(bold=True); _borda(ca); _borda(cb)
    return ln+1

def aba_liquidacao_xl(ws, dados):
    ws.column_dimensions["A"].width = 30
    ws.column_dimensions["B"].width = 50
    h = ws.cell(1,1,"LIQUIDAÇÃO DE SENTENÇA")
    h.font = Font(bold=True,size=13,color=BRANCO)
    h.fill = PatternFill("solid",fgColor=AZUL)
    h.alignment = Alignment(horizontal="center")
    ws.merge_cells("A1:B1")
    n = 3
    for lb, ch in [
        ("Número do processo","numero_processo"),("",""),
        ("Vara do Trabalho","vara_trabalho"),("",""),
        ("Reclamante","reclamante"),("CPF","cpf_reclamante"),
        ("Advogado reclamante","adv_reclamante"),("",""),
        ("1ª Reclamada","reclamada_1"),("CNPJ","cnpj_reclamada_1"),
        ("Advogado reclamada","adv_reclamada_1"),("",""),
        ("Admissão","data_admissao"),("Demissão","data_demissao"),
        ("Ajuizamento","data_ajuizamento"),
    ]:
        if lb == "": n += 1; continue
        n = _campo_xl(ws, n, lb, dados.get(ch))

def aba_pagamentos_xl(ws, ficha):
    if not ficha or "erro" in ficha or not ficha.get("competencias"):
        ws.cell(1,1,"Ficha não localizada ou sem tabelas detectáveis."); return
    rubricas = ficha.get("rubricas",[])
    cabs = ["Competência"]
    for r in rubricas: cabs += [f"{r} Ref", f"{r} Valor"]
    cabs += ["Total Proventos","INSS Base","INSS Desconto","FGTS Base"]
    for ci,t in enumerate(cabs,1):
        c = ws.cell(1,ci,t); _cab(c)
        ws.column_dimensions[get_column_letter(ci)].width = 16
    for li, comp in enumerate(ficha["competencias"],2):
        col=1; ws.cell(li,col,comp.get("competencia","")); col+=1
        for r in rubricas:
            v = comp.get("valores",{}).get(r,{})
            ws.cell(li,col,v.get("referencia","")); ws.cell(li,col+1,v.get("valor",0)); col+=2
        ws.cell(li,col,comp.get("total_proventos",""))
        ws.cell(li,col+1,comp.get("inss_base",""))
        ws.cell(li,col+2,comp.get("inss_desconto",""))
        ws.cell(li,col+3,comp.get("fgts_base",""))
    ws.freeze_panes = "B2"

def aba_ponto_xl(ws, ponto):
    if not ponto or "erro" in ponto or not ponto.get("registros"):
        ws.cell(1,1,"Ponto não localizado ou sem tabelas detectáveis."); return
    regs = ponto["registros"]
    n_max = max((len(r.get("entradas_saidas",[])) for r in regs), default=4)
    n_max = max(n_max,2)
    cabs = ["Data","Dia"]
    for i in range(1,n_max//2+1): cabs += [f"E{i}",f"S{i}"]
    cabs += ["H.Trab.","H.Extra","Observação"]
    for ci,t in enumerate(cabs,1):
        c = ws.cell(1,ci,t); _cab(c)
        ws.column_dimensions[get_column_letter(ci)].width = 12
    for li,reg in enumerate(regs,2):
        ws.cell(li,1,reg.get("data","")); ws.cell(li,2,reg.get("dia_semana",""))
        for ci,hs in enumerate(reg.get("entradas_saidas",[]),3):
            ws.cell(li,ci,hs)
        ultimo = 2+n_max
        ws.cell(li,ultimo+1,reg.get("horas_trabalhadas",""))
        ws.cell(li,ultimo+2,reg.get("horas_extras",""))
        ws.cell(li,ultimo+3,reg.get("observacao",""))
    ws.freeze_panes = "A2"

def gerar_excel(dados, ficha, ponto, inc_pag, inc_pto) -> bytes:
    wb = Workbook()
    ws = wb.active; ws.title = "LIQUIDACAO"
    aba_liquidacao_xl(ws, dados)
    if inc_pag: aba_pagamentos_xl(wb.create_sheet("PAGAMENTOS"), ficha)
    if inc_pto: aba_ponto_xl(wb.create_sheet("PONTO"), ponto)
    buf = io.BytesIO(); wb.save(buf); buf.seek(0)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
# BLOCO 4 — INTERFACE
# ══════════════════════════════════════════════════════════════════════════════

arq = st.file_uploader("📄 PDF do processo", type=["pdf"])
if not arq:
    st.info("Faça o upload do PDF para continuar.")
    st.stop()
st.success(f"{arq.name} ({arq.size/1024/1024:.1f} MB)")

st.subheader("O que extrair?")
c1, c2 = st.columns(2)
with c1:
    f_dados     = st.checkbox("Dados do processo e partes",        value=True)
    f_decisoes  = st.checkbox("Decisões judiciais (dispositivos)", value=True)
    f_criterios = st.checkbox("Critérios de liquidação ⚡IA",      value=True)
    f_alertas   = st.checkbox("Alertas periciais ⚡IA",            value=True,
                               help="Seção 8: riscos, prazos, documentação faltante, PJe-Calc")
with c2:
    f_ficha = st.checkbox("Ficha financeira (holerites)")
    f_ponto = st.checkbox("Espelho de ponto")
st.caption("⚡IA = chamada ao Groq. Todo o resto é processamento local, sem custo.")

st.subheader("Arquivos de saída")
c3, c4 = st.columns(2)
with c3:
    o_word = st.checkbox("Word (.docx)",  value=True)
    o_md   = st.checkbox("Markdown (.md)")
with c4:
    o_xl   = st.checkbox("Excel (.xlsx)")
    if o_xl:
        st.caption("Abas Excel:")
        st.checkbox("LIQUIDACAO (sempre incluída)", value=True, disabled=True)
        aba_pag = st.checkbox("PAGAMENTOS (ficha)", disabled=not f_ficha,
                              help="Marque 'Ficha financeira' acima")
        aba_pto = st.checkbox("PONTO (espelho)",    disabled=not f_ponto,
                              help="Marque 'Espelho de ponto' acima")
    else:
        aba_pag = aba_pto = False

if not (o_word or o_md or o_xl):
    st.warning("Selecione pelo menos um arquivo de saída.")
    st.stop()

if not st.button("⚙️ Processar", type="primary"):
    st.stop()

# ── Processamento ─────────────────────────────────────────────────────────────

pdf_bytes = arq.read()
dados = decisoes_lista = criterios = ficha = ponto = alertas = None

with st.spinner("Lendo PDF..."):
    doc_fitz, capa, txt, toc, npags = ler_pdf(pdf_bytes)
    secs = buscar_secoes(doc_fitz, toc, txt)

secoes_ok = list(secs.keys())
st.write(f"✅ {npags} páginas lidas" +
         (f" — seções: {', '.join(secoes_ok)}" if secoes_ok else " — nenhuma seção localizada por palavra-chave"))

if f_dados:
    with st.spinner("Extraindo dados do processo (regex)..."):
        dados = extrair_dados(txt, capa)
    num = dados.get("numero_processo","?")
    rec = dados.get("reclamante","?")
    emp = dados.get("reclamada_1","?")
    st.write(f"✅ Processo {num} — {rec} × {emp}")

if f_decisoes:
    with st.spinner("Extraindo decisões (texto literal)..."):
        decisoes_lista = extrair_decisoes(secs)
    n = len(decisoes_lista)
    if n:
        tipos = ", ".join(d["tipo"] for d in decisoes_lista)
        st.write(f"✅ {n} decisão(ões): {tipos}")
    else:
        st.write("⚠️ Nenhuma decisão localizada")

if f_criterios:
    dispositivo = secs.get("dispositivo") or secs.get("sentenca") or txt[-30000:]
    with st.spinner("Extraindo critérios de liquidação (⚡IA — 1 chamada)..."):
        criterios = ext_criterios(dispositivo)
    if "erro" in criterios:
        st.error(f"❌ Critérios: {criterios['erro']}")
    else:
        st.write("✅ Critérios extraídos")

if f_alertas:
    # Usa sentença + decisões + despacho de nomeação como fonte
    texto_alertas = "\n\n".join(filter(None, [
        secs.get("sentenca","")[:10000],
        secs.get("acordao","")[:5000],
        secs.get("decisao","")[:5000],
    ])) or txt[-20000:]
    with st.spinner("Gerando alertas periciais (⚡IA — 1 chamada)..."):
        alertas = ext_alertas_periciais(texto_alertas)
    if "erro" in alertas:
        st.error(f"❌ Alertas: {alertas['erro']}")
    else:
        n_alertas = len(alertas.get("alertas", {}).get("pontos_atencao", []))
        st.write(f"✅ Alertas gerados ({n_alertas} ponto(s) de atenção)")

if f_ficha:
    with st.spinner("Extraindo ficha financeira (tabelas PyMuPDF)..."):
        ficha = extrair_ficha(doc_fitz)
    if "erro" in ficha:
        st.warning(f"⚠️ Ficha: {ficha['erro']}")
    else:
        n = len(ficha.get("competencias",[]))
        st.write(f"✅ Ficha: {n} competência(s), {len(ficha.get('rubricas',[]))} rubrica(s)")

if f_ponto:
    with st.spinner("Extraindo espelho de ponto (tabelas PyMuPDF)..."):
        ponto = extrair_ponto(doc_fitz)
    if "erro" in ponto:
        st.warning(f"⚠️ Ponto: {ponto['erro']}")
    else:
        st.write(f"✅ Ponto: {len(ponto.get('registros',[]))} registro(s)")

doc_fitz.close()

# ── Downloads ─────────────────────────────────────────────────────────────────

st.success("✅ Concluído!")
st.subheader("📥 Downloads")
num = (dados or {}).get("numero_processo","processo").replace("-","").replace(".","")

if o_word:
    with st.spinner("Gerando Word..."):
        wb = gerar_word(dados or {}, decisoes_lista or [], criterios or {}, alertas or {})
    st.download_button("📄 Word (.docx)", wb, f"sintese_{num}.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if o_md:
    with st.spinner("Gerando Markdown..."):
        md = gerar_markdown(dados or {}, decisoes_lista or [], criterios or {}, alertas or {})
    st.download_button("📝 Markdown (.md)", md.encode(), f"sintese_{num}.md","text/markdown")

if o_xl:
    with st.spinner("Gerando Excel..."):
        xl = gerar_excel(dados or {}, ficha or {}, ponto or {}, aba_pag, aba_pto)
    st.download_button("📊 Excel (.xlsx)", xl, f"liquidacao_{num}.xlsx",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
