"""Report exporters for SinteseProc."""

import io
import math

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

def gerar_word(dados, decisoes, criterios, alertas=None) -> bytes:
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)

    def h(t, n=1):
        x = doc.add_heading(t, level=n)
        x.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def campo(label, val):
        if not val: return
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(str(val))

    def campo_tabela(tabela, label, val):
        if not val:
            return
        row = tabela.add_row()
        row.cells[0].text = label
        row.cells[1].text = str(val)
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
        row.cells[0].paragraphs[0].runs[0].bold = True

    def texto_multilinha(titulo, texto):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        p.add_run(titulo).bold = True

        linhas = [l.strip() for l in str(texto or "").replace("\r\n", "\n").split("\n")]
        linhas = [l for l in linhas if l]
        if not linhas:
            linhas = ["(não localizado)"]

        for linha in linhas:
            par = doc.add_paragraph()
            par.paragraph_format.left_indent = Inches(0.25)
            par.paragraph_format.first_line_indent = Inches(0)
            par.paragraph_format.space_after = Pt(3)
            par.paragraph_format.line_spacing = 1.05
            par.add_run(linha)

    h("SÍNTESE DE DECISÕES PARA LIQUIDAÇÃO")

    h("1. Identificação do Processo", 2)
    campo("Número",    dados.get("numero_processo"))
    campo("Vara",      dados.get("vara_trabalho"))
    campo("Reclamante",dados.get("reclamante"))
    campo("CPF",       dados.get("cpf_reclamante"))
    campo("Reclamada", dados.get("reclamada_1"))
    campo("CNPJ",      dados.get("cnpj_reclamada_1"))
    campo("Advogado (reclamante)", dados.get("adv_reclamante"))
    campo("Advogado (reclamada)",  dados.get("adv_reclamada_1"))
    campo("Admissão",    dados.get("data_admissao"))
    campo("Demissão",    dados.get("data_demissao"))
    campo("Ajuizamento", dados.get("data_ajuizamento"))
    if dados.get("oabs"):
        campo("OABs encontradas", " | ".join(dados["oabs"]))

    if decisoes:
        h("2. Decisões Judiciais", 2)
        for i, d in enumerate(decisoes, 1):
            h(f"2.{i} {d.get('tipo','Decisão')}", 3)
            meta = doc.add_table(rows=0, cols=2)
            meta.style = "Table Grid"
            meta.columns[0].width = Inches(1.7)
            meta.columns[1].width = Inches(5.4)
            campo_tabela(meta, "Data", d.get("data"))
            campo_tabela(meta, "ID do documento", d.get("id_documento"))
            if d.get("pagina_inicial"):
                pagina_final = d.get("pagina_final") or d.get("pagina_inicial")
                campo_tabela(meta, "Páginas", f"{d.get('pagina_inicial')} a {pagina_final}")
            campo_tabela(meta, "Título de origem", d.get("titulo_origem"))
            campo_tabela(meta, "Resultado", d.get("resultado_reclamante"))
            if d.get("verbas_deferidas"):
                campo_tabela(meta, "Verbas identificadas", ", ".join(d["verbas_deferidas"]))
            texto_multilinha("Dispositivo", d.get("dispositivo",""))

    if criterios and "erro" not in criterios and criterios.get("criterios"):
        h("3. Critérios de Liquidação", 2)
        c = criterios["criterios"]
        per = c.get("periodo_apurado",{})
        if per: campo("Período", f"{per.get('inicio')} a {per.get('fim')}")
        for lb, ch in [
            ("Base salarial","base_salarial"),
            ("Jornada contratual","jornada_contratual"),
            ("Jornada real apurada","jornada_real_apurada"),
            ("Divisor","divisor"),
            ("Adicional HE","adicional_horas_extras"),
            ("Atualização monetária","atualizacao_monetaria"),
            ("Juros","juros"),
            ("INSS empregado","inss_empregado"),
            ("INSS patronal","inss_patronal"),
            ("IR","ir"),
            ("Observações","observacoes"),
        ]:
            campo(lb, c.get(ch))
        if c.get("reflexos"):
            campo("Reflexos", ", ".join(c["reflexos"]))
        fgts = c.get("fgts",{})
        if fgts:
            campo("FGTS", f"{fgts.get('base','')} — {'com' if fgts.get('multa_40') else 'sem'} multa 40%")
        if c.get("exclusoes_expressas"):
            campo("Exclusões expressas", "; ".join(c["exclusoes_expressas"]))
        if c.get("marco_tema9"):
            campo("Marco Tema 9/TST", c["marco_tema9"])

    # Seção 4 — Alertas Periciais (obrigatório conforme instrucoes-analise.md)
    if alertas and "erro" not in alertas and alertas.get("alertas"):
        a = alertas["alertas"]
        h("4. Alertas Periciais", 2)
        p = doc.add_paragraph()
        run = p.add_run("⚠️ Esta seção deve ser lida antes de iniciar o laudo.")
        run.bold = True; run.font.color.rgb = None

        def alerta(label, val):
            if not val: return
            if isinstance(val, list) and not val: return
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{label}: ").bold = True
            p.add_run(", ".join(val) if isinstance(val, list) else str(val))

        campo("Formato do laudo",    a.get("formato_laudo"))
        campo("Prazo",               a.get("prazo_laudo"))
        campo("PJe-Calc exigido",    "Sim" if a.get("pje_calc_exigido") else "Verificar")
        campo("Apuração vincendas",  a.get("vincendas"))
        campo("Risco honorários",    a.get("honorarios_risco"))

        if a.get("documentacao_faltante"):
            h("Documentação faltante", 3)
            for item in a["documentacao_faltante"]:
                doc.add_paragraph(item, style="List Bullet")

        if a.get("pontos_atencao"):
            h("Pontos de atenção", 3)
            for item in a["pontos_atencao"]:
                doc.add_paragraph(item, style="List Bullet")

        if a.get("marcos_temporais_criticos"):
            h("Marcos temporais críticos", 3)
            for item in a["marcos_temporais_criticos"]:
                doc.add_paragraph(item, style="List Bullet")

        campo("Observações", a.get("observacoes_finais"))

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue()


def gerar_markdown(dados, decisoes, criterios, alertas=None) -> str:
    md = ["# SÍNTESE DE DECISÕES PARA LIQUIDAÇÃO\n",
          "## 1. Identificação do Processo\n"]

    for lb, ch in [
        ("Número","numero_processo"),("Vara","vara_trabalho"),
        ("Reclamante","reclamante"),("CPF","cpf_reclamante"),
        ("Reclamada","reclamada_1"),("CNPJ","cnpj_reclamada_1"),
        ("Advogado reclamante","adv_reclamante"),
        ("Advogado reclamada","adv_reclamada_1"),
        ("Admissão","data_admissao"),("Demissão","data_demissao"),
        ("Ajuizamento","data_ajuizamento"),
    ]:
        if dados.get(ch): md.append(f"**{lb}:** {dados[ch]}")
    if dados.get("oabs"):
        md.append(f"**OABs:** {' | '.join(dados['oabs'])}")

    if decisoes:
        md.append("\n## 2. Decisões Judiciais\n")
        for i, d in enumerate(decisoes, 1):
            md.append(f"### 2.{i} {d.get('tipo','Decisão')}")
            md.append(f"**Data:** {d.get('data','')}  |  **Resultado:** {d.get('resultado_reclamante','')}")
            if d.get("id_documento"):
                md.append(f"**ID do documento:** {d['id_documento']}")
            if d.get("pagina_inicial"):
                pagina_final = d.get("pagina_final") or d.get("pagina_inicial")
                md.append(f"**Páginas:** {d.get('pagina_inicial')} a {pagina_final}")
            if d.get("titulo_origem"):
                md.append(f"**Título de origem:** {d['titulo_origem']}")
            if d.get("verbas_deferidas"):
                md.append(f"**Verbas:** {', '.join(d['verbas_deferidas'])}")
            md.append(f"\n**DISPOSITIVO:**\n\n{d.get('dispositivo','')}\n")

    if criterios and "erro" not in criterios and criterios.get("criterios"):
        md.append("\n## 3. Critérios de Liquidação\n")
        c = criterios["criterios"]
        per = c.get("periodo_apurado",{})
        if per: md.append(f"**Período:** {per.get('inicio')} a {per.get('fim')}")
        for lb, ch in [
            ("Base salarial","base_salarial"),("Jornada contratual","jornada_contratual"),
            ("Jornada real","jornada_real_apurada"),("Divisor","divisor"),
            ("Adicional HE","adicional_horas_extras"),("Atualização","atualizacao_monetaria"),
            ("Juros","juros"),("INSS empregado","inss_empregado"),
            ("INSS patronal","inss_patronal"),("IR","ir"),("Observações","observacoes"),
        ]:
            if c.get(ch): md.append(f"**{lb}:** {c[ch]}")
        if c.get("reflexos"): md.append(f"**Reflexos:** {', '.join(c['reflexos'])}")
        fgts = c.get("fgts",{})
        if fgts: md.append(f"**FGTS:** {fgts.get('base','')} — {'com' if fgts.get('multa_40') else 'sem'} multa 40%")
        if c.get("exclusoes_expressas"):
            md.append(f"**Exclusões:** {'; '.join(c['exclusoes_expressas'])}")
        if c.get("marco_tema9"):
            md.append(f"**Marco Tema 9/TST:** {c['marco_tema9']}")

    if alertas and "erro" not in alertas and alertas.get("alertas"):
        a = alertas["alertas"]
        md.append("\n## 4. Alertas Periciais ⚠️\n")
        if a.get("formato_laudo"):   md.append(f"**Formato:** {a['formato_laudo']}")
        if a.get("prazo_laudo"):     md.append(f"**Prazo:** {a['prazo_laudo']}")
        if a.get("pje_calc_exigido") is not None:
            md.append(f"**PJe-Calc:** {'Exigido' if a['pje_calc_exigido'] else 'Verificar'}")
        if a.get("vincendas"):       md.append(f"**Vincendas:** {a['vincendas']}")
        if a.get("honorarios_risco"):md.append(f"**Risco honorários:** {a['honorarios_risco']}")
        if a.get("documentacao_faltante"):
            md.append("\n### Documentação faltante")
            for item in a["documentacao_faltante"]: md.append(f"- {item}")
        if a.get("pontos_atencao"):
            md.append("\n### Pontos de atenção")
            for item in a["pontos_atencao"]: md.append(f"- {item}")
        if a.get("marcos_temporais_criticos"):
            md.append("\n### Marcos temporais críticos")
            for item in a["marcos_temporais_criticos"]: md.append(f"- {item}")
        if a.get("observacoes_finais"): md.append(f"\n**Observações:** {a['observacoes_finais']}")

    return "\n\n".join(md)


AZUL, BRANCO = "1F497D", "FFFFFF"

def _cab(c):
    c.font = Font(bold=True, color=BRANCO, size=11)
    c.fill = PatternFill("solid", fgColor=AZUL)
    c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

def _borda(c):
    b = Side(style="thin")
    c.border = Border(left=b, right=b, top=b, bottom=b)

def _campo_xl(ws, ln, label, val):
    ca = ws.cell(ln,1,label); cb = ws.cell(ln,2,val or "")
    ca.font = Font(bold=True); _borda(ca); _borda(cb)
    return ln+1

def aba_liquidacao_xl(ws, dados):
    ws.column_dimensions["A"].width = 30
    ws.column_dimensions["B"].width = 50
    h = ws.cell(1,1,"LIQUIDAÇÃO DE SENTENÇA")
    h.font = Font(bold=True,size=13,color=BRANCO)
    h.fill = PatternFill("solid",fgColor=AZUL)
    h.alignment = Alignment(horizontal="center")
    ws.merge_cells("A1:B1")
    n = 3
    for lb, ch in [
        ("Número do processo","numero_processo"),("",""),
        ("Vara do Trabalho","vara_trabalho"),("",""),
        ("Reclamante","reclamante"),("CPF","cpf_reclamante"),
        ("Advogado reclamante","adv_reclamante"),("",""),
        ("1ª Reclamada","reclamada_1"),("CNPJ","cnpj_reclamada_1"),
        ("Advogado reclamada","adv_reclamada_1"),("",""),
        ("Admissão","data_admissao"),("Demissão","data_demissao"),
        ("Ajuizamento","data_ajuizamento"),
    ]:
        if lb == "": n += 1; continue
        n = _campo_xl(ws, n, lb, dados.get(ch))

def aba_pagamentos_xl(ws, ficha):
    if not ficha or "erro" in ficha or not ficha.get("competencias"):
        ws.cell(1,1,"Ficha não localizada ou sem tabelas detectáveis."); return
    rubricas = ficha.get("rubricas",[])
    cabs = ["Competência"]
    for r in rubricas: cabs += [f"{r} Ref", f"{r} Valor"]
    cabs += ["Total Proventos","INSS Base","INSS Desconto","FGTS Base"]
    for ci,t in enumerate(cabs,1):
        c = ws.cell(1,ci,t); _cab(c)
        ws.column_dimensions[get_column_letter(ci)].width = 16
    for li, comp in enumerate(ficha["competencias"],2):
        col=1; ws.cell(li,col,comp.get("competencia","")); col+=1
        for r in rubricas:
            v = comp.get("valores",{}).get(r,{})
            ws.cell(li,col,v.get("referencia","")); ws.cell(li,col+1,v.get("valor",0)); col+=2
        ws.cell(li,col,comp.get("total_proventos",""))
        ws.cell(li,col+1,comp.get("inss_base",""))
        ws.cell(li,col+2,comp.get("inss_desconto",""))
        ws.cell(li,col+3,comp.get("fgts_base",""))
    ws.freeze_panes = "B2"

def aba_ponto_xl(ws, ponto):
    if not ponto or "erro" in ponto or not ponto.get("registros"):
        ws.cell(1,1,"Ponto não localizado ou sem tabelas detectáveis."); return
    regs = ponto["registros"]
    n_max = max((len(r.get("entradas_saidas",[])) for r in regs), default=4)
    n_max = max(n_max,2)
    pares = max(1, math.ceil(n_max / 2))
    cabs = ["Data","Dia"]
    for i in range(1, pares + 1): cabs += [f"E{i}",f"S{i}"]
    cabs += ["Col_R1","Col_R2","Col_R3","Col_R4","Observação","Detalhe","H.Trab.","H.Extra"]
    for ci,t in enumerate(cabs,1):
        c = ws.cell(1,ci,t); _cab(c)
        ws.column_dimensions[get_column_letter(ci)].width = 12
    for li,reg in enumerate(regs,2):
        ws.cell(li,1,reg.get("data","")); ws.cell(li,2,reg.get("dia_semana",""))
        for ci,hs in enumerate(reg.get("entradas_saidas",[]),3):
            ws.cell(li,ci,hs)
        ultimo = 2 + (pares * 2)
        for offset in range(1, 5):
            ws.cell(li, ultimo + offset, 0)
        ws.cell(li,ultimo+5,reg.get("observacao",""))
        ws.cell(li,ultimo+6,reg.get("detalhe",""))
        ws.cell(li,ultimo+7,reg.get("horas_trabalhadas",""))
        ws.cell(li,ultimo+8,reg.get("horas_extras",""))
    ws.freeze_panes = "A2"

def gerar_excel(dados, ficha, ponto, inc_pag, inc_pto) -> bytes:
    wb = Workbook()
    ws = wb.active; ws.title = "LIQUIDACAO"
    aba_liquidacao_xl(ws, dados)
    if inc_pag: aba_pagamentos_xl(wb.create_sheet("PAGAMENTOS"), ficha)
    if inc_pto: aba_ponto_xl(wb.create_sheet("PONTO"), ponto)
    buf = io.BytesIO(); wb.save(buf); buf.seek(0)
    return buf.getvalue()
