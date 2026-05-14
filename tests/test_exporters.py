import io
import json

from docx import Document

from sintese.exporters import gerar_json_bytes, gerar_markdown, gerar_word


def test_gerar_word_formata_decisoes_com_tabela_e_dispositivo_multilinha():
    decisoes = [
        {
            "tipo": "Sentença",
            "data": "10/04/2024",
            "id_documento": "abc12345",
            "titulo_origem": "1. 10/04/2024 - Sentença - abc12345",
            "pagina_inicial": 12,
            "pagina_final": 14,
            "resultado_reclamante": "parcialmente procedente",
            "verbas_deferidas": ["horas extras", "FGTS"],
            "dispositivo": "Ante o exposto, julgo procedente.\nCondeno ao pagamento de horas extras.\nIntimem-se.",
            "auditoria_extracao": {
                "ocr_aplicado": True,
                "confianca_minima": "BAIXO",
                "alertas": ["decisão extraída de página OCRizada"],
            },
        }
    ]

    docx_bytes = gerar_word({}, decisoes, {}, {})
    doc = Document(io.BytesIO(docx_bytes))
    textos = [p.text for p in doc.paragraphs if p.text]
    tabelas = [
        cell.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
    ]

    assert "2.1 Sentença" in textos
    assert "Dispositivo" in textos
    assert "Ante o exposto, julgo procedente." in textos
    assert "Condeno ao pagamento de horas extras." in textos
    assert "Intimem-se." in textos
    assert "Data" in tabelas
    assert "10/04/2024" in tabelas
    assert "ID do documento" in tabelas
    assert "abc12345" in tabelas
    assert "Páginas" in tabelas
    assert "12 a 14" in tabelas
    assert "Título de origem" in tabelas
    assert "1. 10/04/2024 - Sentença - abc12345" in tabelas
    assert "Auditoria da extração" in tabelas
    assert "decisão extraída de página OCRizada" in tabelas
    assert "Verbas identificadas" in tabelas
    assert "horas extras, FGTS" in tabelas


def test_gerar_markdown_inclui_auditoria_decisao():
    md = gerar_markdown(
        {},
        [{
            "tipo": "Sentença",
            "data": "10/04/2024",
            "resultado_reclamante": "procedente",
            "dispositivo": "Julgo procedente.",
            "auditoria_extracao": {"alertas": ["decisão contém página de baixa confiança técnica"]},
        }],
        {},
        {},
    )

    assert "Auditoria da extração" in md
    assert "decisão contém página de baixa confiança técnica" in md


def test_gerar_json_bytes_serializa_utf8():
    payload = gerar_json_bytes({"confiança": "BAIXO"})

    assert json.loads(payload.decode("utf-8")) == {"confiança": "BAIXO"}
